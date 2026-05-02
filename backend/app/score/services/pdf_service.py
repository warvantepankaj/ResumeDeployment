import io
from html import escape
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.platypus.flowables import HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.colors import HexColor

from app.score.models.resume_model import ResumeData


def generate_pdf(data: ResumeData) -> bytes:
    """Generate a PDF file from resume data using ReportLab natively."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=HexColor('#111827'),
        alignment=1,  # Center
        spaceAfter=6
    )

    contact_style = ParagraphStyle(
        'ContactStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#6b7280'),
        alignment=1,  # Center
        spaceAfter=12
    )

    section_title_style = ParagraphStyle(
        'SectionTitleStyle',
        parent=styles['Heading2'],
        fontSize=11,
        textColor=HexColor('#4f46e5'),
        spaceAfter=6,
        textTransform='uppercase'
    )

    item_title_style = ParagraphStyle(
        'ItemTitleStyle',
        parent=styles['Heading3'],
        fontSize=10,
        textColor=HexColor('#1f2937'),
        spaceAfter=2
    )

    date_style = ParagraphStyle(
        'DateStyle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#6b7280'),
        spaceAfter=4
    )

    normal_style = ParagraphStyle(
        'NormalStyle',
        parent=styles['Normal'],
        fontSize=9.5,
        textColor=HexColor('#374151'),
        spaceAfter=6,
        leading=14
    )

    bullet_style = ParagraphStyle(
        'BulletStyle',
        parent=styles['Normal'],
        fontSize=9.5,
        textColor=HexColor('#374151'),
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=2,
        leading=14
    )

    story = []
    pi = data.personal_info

    # Header
    name = pi.full_name.upper() if pi.full_name else 'YOUR NAME'
    story.append(Paragraph(escape(name), title_style))

    contact_parts = []
    if pi.email: contact_parts.append(escape(pi.email))
    if pi.phone: contact_parts.append(escape(pi.phone))
    if pi.location: contact_parts.append(escape(pi.location))
    if pi.linkedin: contact_parts.append(escape(pi.linkedin))
    if pi.github: contact_parts.append(escape(pi.github))
    if pi.portfolio: contact_parts.append(escape(pi.portfolio))

    if contact_parts:
        story.append(Paragraph(' | '.join(contact_parts), contact_style))

    story.append(HRFlowable(width="100%", thickness=1, color=HexColor('#4f46e5'), spaceAfter=10, dash=None))

    if pi.summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", section_title_style))
        story.append(Paragraph(escape(pi.summary), normal_style))
        story.append(Spacer(1, 10))

    if data.experience:
        story.append(Paragraph("EXPERIENCE", section_title_style))
        for exp in data.experience:
            end = "Present" if exp.current else escape(exp.end_date)
            location_part = f" | {escape(exp.location)}" if exp.location else ""
            title_text = f"<b>{escape(exp.position)}</b> | {escape(exp.company)}{location_part}"
            story.append(Paragraph(title_text, item_title_style))
            story.append(Paragraph(f"{escape(exp.start_date)} – {end}", date_style))

            if exp.description:
                story.append(Paragraph(escape(exp.description), normal_style))

            for hl in exp.highlights:
                # Add bullet point character
                story.append(Paragraph(f"• {escape(hl)}", bullet_style))
        story.append(Spacer(1, 10))

    if data.education:
        story.append(Paragraph("EDUCATION", section_title_style))
        for edu in data.education:
            gpa_part = f" | GPA: {escape(edu.gpa)}" if edu.gpa else ""
            title_text = f"<b>{escape(edu.degree)} in {escape(edu.field_of_study)}</b> | {escape(edu.institution)}"
            story.append(Paragraph(title_text, item_title_style))
            story.append(Paragraph(f"{escape(edu.start_date)} – {escape(edu.end_date)}{gpa_part}", date_style))
            if edu.description:
                story.append(Paragraph(escape(edu.description), normal_style))
        story.append(Spacer(1, 10))

    if data.skills:
        story.append(Paragraph("SKILLS", section_title_style))
        skills_text = ", ".join(escape(s) for s in data.skills)
        story.append(Paragraph(skills_text, normal_style))
        story.append(Spacer(1, 10))

    if data.projects:
        story.append(Paragraph("PROJECTS", section_title_style))
        for proj in data.projects:
            story.append(Paragraph(f"<b>{escape(proj.name)}</b>", item_title_style))

            date_part = ""
            if proj.start_date:
                date_part = escape(proj.start_date)
                if proj.end_date:
                    date_part += f" – {escape(proj.end_date)}"
            if date_part:
                story.append(Paragraph(date_part, date_style))

            if proj.technologies:
                tech_text = f"<i>{escape(', '.join(proj.technologies))}</i>"
                story.append(Paragraph(tech_text, date_style))

            if proj.description:
                story.append(Paragraph(escape(proj.description), normal_style))
        story.append(Spacer(1, 10))

    if data.certifications:
        story.append(Paragraph("CERTIFICATIONS", section_title_style))
        for cert in data.certifications:
            issuer_part = f" – {escape(cert.issuer)}" if cert.issuer else ""
            date_part = f" ({escape(cert.date)})" if cert.date else ""
            cert_text = f"<b>{escape(cert.name)}</b>{issuer_part} {date_part}"
            story.append(Paragraph(cert_text, normal_style))
        story.append(Spacer(1, 10))

    if data.languages:
        story.append(Paragraph("LANGUAGES", section_title_style))
        story.append(Paragraph(" • ".join(escape(l) for l in data.languages), normal_style))
        story.append(Spacer(1, 10))

    if data.interests:
        story.append(Paragraph("INTERESTS", section_title_style))
        story.append(Paragraph(" • ".join(escape(i) for i in data.interests), normal_style))
        story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()


def generate_docx(data: ResumeData) -> bytes:
    """Generate a DOCX file from resume data."""
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    pi = data.personal_info

    # Name
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(pi.full_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(18)
    name_run.font.color.rgb = RGBColor(44, 62, 80)

    # Contact info
    contact_parts = []
    if pi.email:
        contact_parts.append(pi.email)
    if pi.phone:
        contact_parts.append(pi.phone)
    if pi.location:
        contact_parts.append(pi.location)
    if pi.linkedin:
        contact_parts.append(pi.linkedin)
    if pi.github:
        contact_parts.append(pi.github)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_run = contact_para.add_run(" | ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = RGBColor(100, 100, 100)

    # Summary
    if pi.summary:
        _add_section_header(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(pi.summary)
        p.style.font.size = Pt(10)

    # Experience
    if data.experience:
        _add_section_header(doc, "EXPERIENCE")
        for exp in data.experience:
            p = doc.add_paragraph()
            title_run = p.add_run(f"{exp.position}")
            title_run.bold = True
            title_run.font.size = Pt(11)
            p.add_run(f"  |  {exp.company}")
            if exp.location:
                p.add_run(f"  |  {exp.location}")

            date_para = doc.add_paragraph()
            date_run = date_para.add_run(
                f"{exp.start_date} - {'Present' if exp.current else exp.end_date}"
            )
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(100, 100, 100)

            if exp.description:
                doc.add_paragraph(exp.description, style="List Bullet")
            for h in exp.highlights:
                doc.add_paragraph(h, style="List Bullet")

    # Education
    if data.education:
        _add_section_header(doc, "EDUCATION")
        for edu in data.education:
            p = doc.add_paragraph()
            deg_run = p.add_run(f"{edu.degree} in {edu.field_of_study}")
            deg_run.bold = True
            deg_run.font.size = Pt(11)
            p.add_run(f"  |  {edu.institution}")

            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"{edu.start_date} - {edu.end_date}")
            date_run.font.size = Pt(9)
            date_run.font.color.rgb = RGBColor(100, 100, 100)
            if edu.gpa:
                date_para.add_run(f"  |  GPA: {edu.gpa}")

    # Skills
    if data.skills:
        _add_section_header(doc, "SKILLS")
        doc.add_paragraph(", ".join(data.skills))

    # Projects
    if data.projects:
        _add_section_header(doc, "PROJECTS")
        for proj in data.projects:
            p = doc.add_paragraph()
            proj_run = p.add_run(proj.name)
            proj_run.bold = True
            proj_run.font.size = Pt(11)
            if proj.technologies:
                p.add_run(f"  |  {', '.join(proj.technologies)}")
            if proj.description:
                doc.add_paragraph(proj.description, style="List Bullet")

    # Certifications
    if data.certifications:
        _add_section_header(doc, "CERTIFICATIONS")
        for cert in data.certifications:
            doc.add_paragraph(f"{cert.name} - {cert.issuer} ({cert.date})")

    # Languages
    if data.languages:
        _add_section_header(doc, "LANGUAGES")
        doc.add_paragraph(", ".join(data.languages))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_section_header(doc: Document, title: str):
    """Add a styled section header with a bottom border."""
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    p.space_after = Pt(4)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(44, 62, 80)

    # Add a thin line
    border_p = doc.add_paragraph()
    border_p.space_before = Pt(0)
    border_p.space_after = Pt(6)
    border_run = border_p.add_run("─" * 80)
    border_run.font.size = Pt(6)
    border_run.font.color.rgb = RGBColor(189, 195, 199)
